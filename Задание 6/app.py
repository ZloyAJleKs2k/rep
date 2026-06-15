# # EMAIL
# import sys
# import re
# import os
# import requests
# import tkinter as tk
# from tkinter import ttk, messagebox
# from docx import Document
#
#
# class EmailValidationApp(tk.Tk):
#     def __init__(self):
#         super().__init__()
#         self.current_input = ""
#         self.current_result = ""
#         self.test_index = 0
#         self.api_url = "http://localhost:4444/TransferSimulator/email"
#         self.init_ui()
#
#     def init_ui(self):
#         self.title("Валидация Email - Бургер плюс")
#         self.minsize(600, 250)
#         self.geometry("650x280")
#
#         # Главный контейнер с отступами
#         main_frame = ttk.Frame(self, padding=20)
#         main_frame.pack(fill=tk.BOTH, expand=True)
#
#         # Первая строка: Получить данные
#         row1 = ttk.Frame(main_frame)
#         row1.pack(fill=tk.X, pady=(0, 15))
#
#         self.btn_get_data = ttk.Button(
#             row1,
#             text="Получить данные",
#             width=25,
#             command=self.get_data
#         )
#         self.btn_get_data.pack(side=tk.LEFT, padx=(0, 10))
#
#         self.lbl_data = tk.Label(
#             row1,
#             text="  ",
#             relief=tk.SOLID,
#             borderwidth=1,
#             anchor=tk.W,
#             padx=8,
#             pady=8,
#             bg="white"
#         )
#         self.lbl_data.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8)
#
#         # Вторая строка: Отправить результат
#         row2 = ttk.Frame(main_frame)
#         row2.pack(fill=tk.X)
#
#         self.btn_send_result = ttk.Button(
#             row2,
#             text="Отправить результат теста",
#             width=25,
#             command=self.send_result
#         )
#         self.btn_send_result.pack(side=tk.LEFT, padx=(0, 10))
#
#         self.lbl_result = tk.Label(
#             row2,
#             text="  ",
#             relief=tk.SOLID,
#             borderwidth=1,
#             anchor=tk.W,
#             padx=8,
#             pady=8,
#             bg="white"
#         )
#         self.lbl_result.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8)
#
#     def get_data(self):
#         try:
#             response = requests.get(self.api_url, timeout=5)
#
#             if response.status_code == 500:
#                 messagebox.showerror(
#                     "Ошибка сервера",
#                     "HttpStatusCode: 500 Internal Server Error.\n"
#                     "Немедленно обратитесь к главному эксперту для фиксации проблемы."
#                 )
#                 return
#
#             response.raise_for_status()
#             data = response.json()
#
#             email = data.get("value", "").strip()
#
#             if not email:
#                 messagebox.showwarning(
#                     "Предупреждение",
#                     "Получен пустой ответ от эмулятора.\n"
#                     "Проверьте корректность работы TransferSimulator."
#                 )
#                 return
#
#             self.current_input = email
#             self.lbl_data.config(text=self.current_input, fg="black")
#             self.lbl_result.config(text="  ", fg="black", bg="white")
#
#         except requests.exceptions.ConnectionError:
#             messagebox.showerror(
#                 "Ошибка подключения",
#                 "Не удалось подключиться к эмулятору.\n"
#                 "Убедитесь, что TransferSimulator.exe запущен."
#             )
#         except requests.exceptions.Timeout:
#             messagebox.showerror(
#                 "Ошибка",
#                 "Превышено время ожидания ответа от эмулятора."
#             )
#         except requests.exceptions.RequestException as e:
#             messagebox.showerror(
#                 "Ошибка запроса",
#                 f"Произошла ошибка при получении данных:\n{str(e)}"
#             )
#         except ValueError:
#             messagebox.showerror(
#                 "Ошибка парсинга",
#                 "Ответ сервера не является корректным JSON."
#             )
#         except Exception as e:
#             messagebox.showerror(
#                 "Непредвиденная ошибка",
#                 f"Произошла ошибка:\n{str(e)}"
#             )
#
#     def validate_email(self, email):
#         """
#         Валидация email на 2 критерия.
#         """
#         email = email.strip()
#
#         if not email:
#             return False, "Email пустой"
#
#         # Критерий 1: Проверка на наличие разделителей (пробел, ;, ,)
#         if ' ' in email or ';' in email or ',' in email:
#             return False, "Email содержит недопустимые разделители"
#
#         # Проверка на количество символов @
#         if email.count('@') != 1:
#             return False, "Email должен содержать ровно один символ @"
#
#         local_part, domain = email.split('@')
#
#         # Проверка локальной части
#         if not local_part:
#             return False, "Локальная часть email пуста"
#
#         # Критерий 2: Формат локальной части
#         if not re.match(r'^[a-zA-Z0-9._-]+$', local_part):
#             return False, "Локальная часть содержит недопустимые символы"
#
#         # Проверка домена
#         if not domain or '.' not in domain:
#             return False, "Некорректный домен"
#
#         return True, "Email корректен"
#
#     def send_result(self):
#         try:
#             if not self.current_input:
#                 messagebox.showwarning(
#                     "Предупреждение",
#                     "Сначала нажмите 'Получить данные'."
#                 )
#                 return
#
#             is_valid, message = self.validate_email(self.current_input)
#
#             if is_valid:
#                 result_text = "Успешно"
#                 self.lbl_result.config(text=message, fg="green", bg="white")
#             else:
#                 result_text = "Не успешно"
#                 self.lbl_result.config(text=message, fg="red", bg="white")
#
#             self.current_result = result_text
#
#             self.write_to_testcase()
#
#             self.test_index += 1
#
#             messagebox.showinfo(
#                 "Результат сохранён",
#                 f"Результат записан в ТестКейс.docx\n\n"
#                 f"Тест {self.test_index}: {self.current_result}\n"
#                 f"Email: {self.current_input}"
#             )
#
#         except Exception as e:
#             messagebox.showerror(
#                 "Ошибка",
#                 f"Произошла ошибка при отправке результата:\n{str(e)}"
#             )
#
#     def write_to_testcase(self):
#         try:
#             script_dir = os.path.dirname(os.path.abspath(__file__))
#             docx_path = os.path.join(script_dir, "ТестКейс.docx")
#
#             if not os.path.exists(docx_path):
#                 messagebox.showwarning(
#                     "Предупреждение",
#                     f"Файл не найден: {docx_path}\n"
#                     f"Создайте документ ТестКейс.docx с таблицей."
#                 )
#                 return
#
#             doc = Document(docx_path)
#
#             if not doc.tables:
#                 messagebox.showwarning(
#                     "Предупреждение",
#                     "Документ не содержит таблиц.\n"
#                     "Создайте таблицу со столбцами: Действие, Ожидаемый результат, Результат."
#                 )
#                 return
#
#             table = doc.tables[0]
#             row_index = self.test_index + 2
#
#             while row_index >= len(table.rows):
#                 table.add_row()
#
#             if row_index < len(table.rows):
#                 is_valid, message = self.validate_email(self.current_input)
#                 action = "Ввод некорректного email" if not is_valid else "Ввод корректного email"
#
#                 table.rows[row_index].cells[0].text = action
#                 table.rows[row_index].cells[1].text = message
#                 table.rows[row_index].cells[2].text = self.current_result
#
#             doc.save(docx_path)
#
#         except Exception as e:
#             messagebox.showerror(
#                 "Ошибка записи",
#                 f"Не удалось записать результат в документ:\n{str(e)}"
#             )
#
#
# def main():
#     try:
#         app = EmailValidationApp()
#         app.mainloop()
#     except Exception as e:
#         print(f"Критическая ошибка приложения: {e}")
#
#
# if __name__ == "__main__":
#     main()

# Phone

import sys
import re
import os
import requests
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document


class PhoneValidationApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.current_input = ""
        self.current_result = ""
        self.test_index = 0
        self.api_url = "http://localhost:4444/TransferSimulator/mobilePhone"
        self.init_ui()

    def init_ui(self):
        self.title("Валидация телефона - Бургер плюс")
        self.minsize(600, 250)
        self.geometry("650x280")

        main_frame = ttk.Frame(self, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        row1 = ttk.Frame(main_frame)
        row1.pack(fill=tk.X, pady=(0, 15))

        self.btn_get_data = ttk.Button(
            row1,
            text="Получить данные",
            width=25,
            command=self.get_data
        )
        self.btn_get_data.pack(side=tk.LEFT, padx=(0, 10))

        self.lbl_data = tk.Label(
            row1,
            text="  ",
            relief=tk.SOLID,
            borderwidth=1,
            anchor=tk.W,
            padx=8,
            pady=8,
            bg="white"
        )
        self.lbl_data.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8)

        row2 = ttk.Frame(main_frame)
        row2.pack(fill=tk.X)

        self.btn_send_result = ttk.Button(
            row2,
            text="Отправить результат теста",
            width=25,
            command=self.send_result
        )
        self.btn_send_result.pack(side=tk.LEFT, padx=(0, 10))

        self.lbl_result = tk.Label(
            row2,
            text="  ",
            relief=tk.SOLID,
            borderwidth=1,
            anchor=tk.W,
            padx=8,
            pady=8,
            bg="white"
        )
        self.lbl_result.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8)

    def get_data(self):
        try:
            response = requests.get(self.api_url, timeout=5)

            if response.status_code == 500:
                messagebox.showerror(
                    "Ошибка сервера",
                    "HttpStatusCode: 500 Internal Server Error.\n"
                    "Немедленно обратитесь к главному эксперту для фиксации проблемы."
                )
                return

            response.raise_for_status()
            data = response.json()

            phone = data.get("value", "").strip()

            if not phone:
                messagebox.showwarning(
                    "Предупреждение",
                    "Получен пустой ответ от эмулятора.\n"
                    "Проверьте корректность работы TransferSimulator."
                )
                return

            self.current_input = phone
            self.lbl_data.config(text=self.current_input, fg="black")
            self.lbl_result.config(text="  ", fg="black", bg="white")

        except requests.exceptions.ConnectionError:
            messagebox.showerror(
                "Ошибка подключения",
                "Не удалось подключиться к эмулятору.\n"
                "Убедитесь, что TransferSimulator.exe запущен."
            )
        except requests.exceptions.Timeout:
            messagebox.showerror(
                "Ошибка",
                "Превышено время ожидания ответа от эмулятора."
            )
        except requests.exceptions.RequestException as e:
            messagebox.showerror(
                "Ошибка запроса",
                f"Произошла ошибка при получении данных:\n{str(e)}"
            )
        except ValueError:
            messagebox.showerror(
                "Ошибка парсинга",
                "Ответ сервера не является корректным JSON."
            )
        except Exception as e:
            messagebox.showerror(
                "Непредвиденная ошибка",
                f"Произошла ошибка:\n{str(e)}"
            )

    def validate_phone(self, phone):
        """
        Валидация телефона на 2 критерия.
        """
        phone = phone.strip()

        if not phone:
            return False, "Номер телефона пустой"

        # Критерий 1: Проверка на наличие букв или недопустимых символов
        if not re.match(r'^[\d\s\-\+\(\)]+$', phone):
            return False, "Номер содержит недопустимые символы"

        # Проверка, что номер не начинается и не заканчивается на спецсимволы
        if phone.startswith(('-', '(', ')')):
            return False, "Номер не должен начинаться с недопустимого символа"

        if phone.endswith(('-', '(', ')', '+', ' ')):
            return False, "Номер не должен заканчиваться на спецсимвол"

        # Проверка, что плюс может быть только в начале
        if '+' in phone:
            if phone[0] != '+':
                return False, "Символ + может быть только в начале номера"
            if phone.count('+') > 1:
                return False, "Символ + может встречаться только один раз"

        # Проверка парности скобок
        if phone.count('(') != phone.count(')'):
            return False, "Скобки должны быть парными"

        # Критерий 2: Длина номера (10-11 цифр)
        digits = re.sub(r'\D', '', phone)
        if len(digits) < 10 or len(digits) > 11:
            return False, "Номер должен содержать 10-11 цифр"

        # Проверка, что после извлечения цифр номер начинается с 7 или 8 (для РФ)
        if not digits.startswith(('7', '8')):
            return False, "Номер должен начинаться с 7 или 8"

        return True, "Номер телефона корректен"

    def send_result(self):
        try:
            if not self.current_input:
                messagebox.showwarning(
                    "Предупреждение",
                    "Сначала нажмите 'Получить данные'."
                )
                return

            is_valid, message = self.validate_phone(self.current_input)

            if is_valid:
                result_text = "Успешно"
                self.lbl_result.config(text=message, fg="green", bg="white")
            else:
                result_text = "Не успешно"
                self.lbl_result.config(text=message, fg="red", bg="white")

            self.current_result = result_text

            self.write_to_testcase()

            self.test_index += 1

            messagebox.showinfo(
                "Результат сохранён",
                f"Результат записан в ТестКейс.docx\n\n"
                f"Тест {self.test_index}: {self.current_result}\n"
                f"Телефон: {self.current_input}"
            )

        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Произошла ошибка при отправке результата:\n{str(e)}"
            )

    def write_to_testcase(self):
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            docx_path = os.path.join(script_dir, "ТестКейс.docx")

            if not os.path.exists(docx_path):
                messagebox.showwarning(
                    "Предупреждение",
                    f"Файл не найден: {docx_path}\n"
                    f"Создайте документ ТестКейс.docx с таблицей."
                )
                return

            doc = Document(docx_path)

            if not doc.tables:
                messagebox.showwarning(
                    "Предупреждение",
                    "Документ не содержит таблиц.\n"
                    "Создайте таблицу со столбцами: Действие, Ожидаемый результат, Результат."
                )
                return

            table = doc.tables[0]
            row_index = self.test_index + 2

            while row_index >= len(table.rows):
                table.add_row()

            if row_index < len(table.rows):
                is_valid, message = self.validate_phone(self.current_input)
                action = "Ввод некорректного номера" if not is_valid else "Ввод корректного номера"

                table.rows[row_index].cells[0].text = action
                table.rows[row_index].cells[1].text = message
                table.rows[row_index].cells[2].text = self.current_result

            doc.save(docx_path)

        except Exception as e:
            messagebox.showerror(
                "Ошибка записи",
                f"Не удалось записать результат в документ:\n{str(e)}"
            )


def main():
    try:
        app = PhoneValidationApp()
        app.mainloop()
    except Exception as e:
        print(f"Критическая ошибка приложения: {e}")


if __name__ == "__main__":
    main()

# # Password
#
# import sys
# import re
# import os
# import requests
# import tkinter as tk
# from tkinter import ttk, messagebox
# from docx import Document
#
#
# class PassportValidationApp(tk.Tk):
#     def __init__(self):
#         super().__init__()
#         self.current_input = ""
#         self.current_result = ""
#         self.test_index = 0
#         self.api_url = "http://localhost:4444/TransferSimulator/identityCard"
#         self.init_ui()
#
#     def init_ui(self):
#         self.title("Валидация паспорта - Бургер плюс")
#         self.minsize(600, 250)
#         self.geometry("650x280")
#
#         main_frame = ttk.Frame(self, padding=20)
#         main_frame.pack(fill=tk.BOTH, expand=True)
#
#         row1 = ttk.Frame(main_frame)
#         row1.pack(fill=tk.X, pady=(0, 15))
#
#         self.btn_get_data = ttk.Button(
#             row1,
#             text="Получить данные",
#             width=25,
#             command=self.get_data
#         )
#         self.btn_get_data.pack(side=tk.LEFT, padx=(0, 10))
#
#         self.lbl_data = tk.Label(
#             row1,
#             text="  ",
#             relief=tk.SOLID,
#             borderwidth=1,
#             anchor=tk.W,
#             padx=8,
#             pady=8,
#             bg="white"
#         )
#         self.lbl_data.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8)
#
#         row2 = ttk.Frame(main_frame)
#         row2.pack(fill=tk.X)
#
#         self.btn_send_result = ttk.Button(
#             row2,
#             text="Отправить результат теста",
#             width=25,
#             command=self.send_result
#         )
#         self.btn_send_result.pack(side=tk.LEFT, padx=(0, 10))
#
#         self.lbl_result = tk.Label(
#             row2,
#             text="  ",
#             relief=tk.SOLID,
#             borderwidth=1,
#             anchor=tk.W,
#             padx=8,
#             pady=8,
#             bg="white"
#         )
#         self.lbl_result.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8)
#
#     def get_data(self):
#         try:
#             response = requests.get(self.api_url, timeout=5)
#
#             if response.status_code == 500:
#                 messagebox.showerror(
#                     "Ошибка сервера",
#                     "HttpStatusCode: 500 Internal Server Error.\n"
#                     "Немедленно обратитесь к главному эксперту для фиксации проблемы."
#                 )
#                 return
#
#             response.raise_for_status()
#             data = response.json()
#
#             passport = data.get("value", "").strip()
#
#             if not passport:
#                 messagebox.showwarning(
#                     "Предупреждение",
#                     "Получен пустой ответ от эмулятора.\n"
#                     "Проверьте корректность работы TransferSimulator."
#                 )
#                 return
#
#             self.current_input = passport
#             self.lbl_data.config(text=self.current_input, fg="black")
#             self.lbl_result.config(text="  ", fg="black", bg="white")
#
#         except requests.exceptions.ConnectionError:
#             messagebox.showerror(
#                 "Ошибка подключения",
#                 "Не удалось подключиться к эмулятору.\n"
#                 "Убедитесь, что TransferSimulator.exe запущен."
#             )
#         except requests.exceptions.Timeout:
#             messagebox.showerror(
#                 "Ошибка",
#                 "Превышено время ожидания ответа от эмулятора."
#             )
#         except requests.exceptions.RequestException as e:
#             messagebox.showerror(
#                 "Ошибка запроса",
#                 f"Произошла ошибка при получении данных:\n{str(e)}"
#             )
#         except ValueError:
#             messagebox.showerror(
#                 "Ошибка парсинга",
#                 "Ответ сервера не является корректным JSON."
#             )
#         except Exception as e:
#             messagebox.showerror(
#                 "Непредвиденная ошибка",
#                 f"Произошла ошибка:\n{str(e)}"
#             )
#
#     def validate_passport(self, passport):
#         """
#         Валидация паспорта на 2 критерия.
#         Формат: "1234 567890" или "1234567890"
#         """
#         passport = passport.strip()
#
#         if not passport:
#             return False, "Паспортные данные пусты"
#
#         # Извлекаем только цифры
#         digits = re.sub(r'\D', '', passport)
#
#         # Критерий 1: Общая длина должна быть 10 цифр
#         if len(digits) != 10:
#             return False, "Паспорт должен содержать 10 цифр"
#
#         # Разделяем на серию (4) и номер (6)
#         series = digits[:4]
#         number = digits[4:]
#
#         # Критерий 2: Серия — 4 цифры, номер — 6 цифр
#         if not re.match(r'^\d{4}$', series):
#             return False, "Серия паспорта должна содержать 4 цифры"
#
#         if not re.match(r'^\d{6}$', number):
#             return False, "Номер паспорта должен содержать 6 цифр"
#
#         return True, "Паспортные данные корректны"
#
#     def send_result(self):
#         try:
#             if not self.current_input:
#                 messagebox.showwarning(
#                     "Предупреждение",
#                     "Сначала нажмите 'Получить данные'."
#                 )
#                 return
#
#             is_valid, message = self.validate_passport(self.current_input)
#
#             if is_valid:
#                 result_text = "Успешно"
#                 self.lbl_result.config(text=message, fg="green", bg="white")
#             else:
#                 result_text = "Не успешно"
#                 self.lbl_result.config(text=message, fg="red", bg="white")
#
#             self.current_result = result_text
#
#             self.write_to_testcase()
#
#             self.test_index += 1
#
#             messagebox.showinfo(
#                 "Результат сохранён",
#                 f"Результат записан в ТестКейс.docx\n\n"
#                 f"Тест {self.test_index}: {self.current_result}\n"
#                 f"Паспорт: {self.current_input}"
#             )
#
#         except Exception as e:
#             messagebox.showerror(
#                 "Ошибка",
#                 f"Произошла ошибка при отправке результата:\n{str(e)}"
#             )
#
#     def write_to_testcase(self):
#         try:
#             script_dir = os.path.dirname(os.path.abspath(__file__))
#             docx_path = os.path.join(script_dir, "ТестКейс.docx")
#
#             if not os.path.exists(docx_path):
#                 messagebox.showwarning(
#                     "Предупреждение",
#                     f"Файл не найден: {docx_path}\n"
#                     f"Создайте документ ТестКейс.docx с таблицей."
#                 )
#                 return
#
#             doc = Document(docx_path)
#
#             if not doc.tables:
#                 messagebox.showwarning(
#                     "Предупреждение",
#                     "Документ не содержит таблиц.\n"
#                     "Создайте таблицу со столбцами: Действие, Ожидаемый результат, Результат."
#                 )
#                 return
#
#             table = doc.tables[0]
#             row_index = self.test_index + 2
#
#             while row_index >= len(table.rows):
#                 table.add_row()
#
#             if row_index < len(table.rows):
#                 is_valid, message = self.validate_passport(self.current_input)
#                 action = "Ввод некорректных паспортных данных" if not is_valid else "Ввод корректных паспортных данных"
#
#                 table.rows[row_index].cells[0].text = action
#                 table.rows[row_index].cells[1].text = message
#                 table.rows[row_index].cells[2].text = self.current_result
#
#             doc.save(docx_path)
#
#         except Exception as e:
#             messagebox.showerror(
#                 "Ошибка записи",
#                 f"Не удалось записать результат в документ:\n{str(e)}"
#             )
#
#
# def main():
#     try:
#         app = PassportValidationApp()
#         app.mainloop()
#     except Exception as e:
#         print(f"Критическая ошибка приложения: {e}")
#
#
# if __name__ == "__main__":
#     main()